import os
import logging
import subprocess
import json
import base64
import hashlib
import re
import tempfile
import concurrent.futures
import zipfile
from functools import partial
import importlib.util

import httpx
import fitz  # PyMuPDF
from docx import Document
from google import genai
from google.genai import types

try:
    from .layout_index import build_pdf_layout_map as _build_pdf_layout_map
except ImportError:
    _layout_spec = importlib.util.spec_from_file_location(
        "lexanchor_layout_index",
        os.path.join(os.path.dirname(__file__), "layout_index.py"),
    )
    if _layout_spec is None or _layout_spec.loader is None:
        raise
    _layout_module = importlib.util.module_from_spec(_layout_spec)
    _layout_spec.loader.exec_module(_layout_module)
    _build_pdf_layout_map = _layout_module.build_pdf_layout_map

logger = logging.getLogger("legal_redaction_tool.extractor")

def _ocr_single_page_worker(input_file, page_num, vision_provider, vision_model, vision_base_url):
    """Worker function to OCR a single page from a PDF file."""
    img_path = None
    try:
        with fitz.open(input_file) as doc:
            page = doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(1, 1))
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                img_path = tmp.name
            pix.save(img_path)
            
            page_text, meta = extract_text_from_image(
                img_path,
                vision_provider=vision_provider,
                vision_model=vision_model,
                vision_base_url=vision_base_url,
            )
            
            ocr_lines = meta.get("ocr_lines", [])
            valid_lines = []
            for line in ocr_lines:
                bbox = line.get("bbox")
                if bbox and len(bbox) == 4 and bbox[0] < bbox[2] and bbox[1] < bbox[3]:
                    valid_lines.append(line)
            
            return page_num, page_text, valid_lines, None
    except Exception as e:
        return page_num, "", [], str(e)
    finally:
        if img_path and os.path.exists(img_path):
            try: os.remove(img_path)
            except: pass

def extract_text_from_pdf(
    input_file,
    *,
    vision_provider: str = "ollama",
    vision_model: str = "gemma4:e4b",
    vision_base_url: str = "http://127.0.0.1:11434",
    checkpoint_path: str | None = None,
):
    text = ""
    pages_ocr_data = {}
    with fitz.open(input_file) as doc:
        for page in doc:
            text += page.get_text()
            
    if not text.strip():
        logger.info(f"No text found in {input_file}. Attempting Concurrent OCR fallback...")
        checkpoint = _load_pdf_ocr_checkpoint(checkpoint_path) or {}
        page_texts = dict(checkpoint.get("page_texts") or {})
        completed_pages = {int(item) for item in (checkpoint.get("completed_pages") or [])}
        failed_pages = {int(item) for item in (checkpoint.get("failed_pages") or [])}
        pages_ocr_data = dict(checkpoint.get("pages_ocr_data") or {})
        
        with fitz.open(input_file) as doc:
            total_pages = len(doc)
            max_pages = int(os.getenv("MAX_OCR_PAGES", "0"))
            pages_to_process = []
            
            for page_num in range(total_pages):
                if max_pages > 0 and page_num >= max_pages: break
                if page_num in completed_pages and str(page_num) in page_texts: continue
                if page_num in failed_pages:
                    logger.warning(f"Skipping previously FAILED page {page_num + 1}/{total_pages}")
                    continue
                pages_to_process.append(page_num)

            if pages_to_process:
                max_workers = min(os.cpu_count() or 4, len(pages_to_process))
                logger.info(f"Dispatching {len(pages_to_process)} pages to {max_workers} processes for OCR...")
                
                with concurrent.futures.ProcessPoolExecutor(max_workers=max_workers) as executor:
                    worker_func = partial(
                        _ocr_single_page_worker, 
                        input_file, 
                        vision_provider=vision_provider,
                        vision_model=vision_model,
                        vision_base_url=vision_base_url
                    )
                    
                    future_to_page = {executor.submit(worker_func, p): p for p in pages_to_process}
                    
                    for future in concurrent.futures.as_completed(future_to_page):
                        p_num, p_text, p_ocr, error = future.result()
                        p_key = str(p_num)
                        if error:
                            logger.error(f"Failed to process page {p_num + 1}: {error}")
                            failed_pages.add(p_num)
                        else:
                            page_texts[p_key] = p_text
                            pages_ocr_data[p_key] = p_ocr
                            completed_pages.add(p_num)
                            logger.info(f"Page {p_num + 1} OCR complete.")
                        
                        # Incremental atomic checkpoint
                        _save_pdf_ocr_checkpoint(checkpoint_path, {
                            "input_file": os.path.abspath(input_file),
                            "total_pages": total_pages,
                            "completed_pages": sorted(list(completed_pages)),
                            "failed_pages": sorted(list(failed_pages)),
                            "page_texts": page_texts,
                            "pages_ocr_data": pages_ocr_data,
                        })

        ordered_pages = []
        for pk in sorted(page_texts.keys(), key=int):
            ordered_pages.append(f"\n--- Page {int(pk) + 1} ---\n{page_texts[pk]}\n")
        text = "".join(ordered_pages)
        
    return text, pages_ocr_data


def build_pdf_layout_map(input_file, extracted_text: str | None = None, pages_ocr_data: dict | None = None):
    return _build_pdf_layout_map(input_file, extracted_text=extracted_text, pages_ocr_data=pages_ocr_data)


def _build_pdf_checkpoint_path(projects_root: str, input_file: str) -> str:
    reports_dir = os.path.join(projects_root, "reports", "ocr_checkpoints")
    os.makedirs(reports_dir, exist_ok=True)
    digest = hashlib.sha1(os.path.abspath(str(input_file or "")).encode("utf-8")).hexdigest()[:16]
    base = os.path.splitext(os.path.basename(str(input_file or "document.pdf")))[0] or "document"
    return os.path.join(reports_dir, f"{base}.{digest}.json")


def _load_pdf_ocr_checkpoint(checkpoint_path: str | None) -> dict:
    if not checkpoint_path or not os.path.exists(checkpoint_path):
        return {}
    try:
        with open(checkpoint_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception as exc:
        logger.warning("Failed to load OCR checkpoint %s: %s", checkpoint_path, exc)
        return {}


def _save_pdf_ocr_checkpoint(checkpoint_path: str | None, payload: dict) -> None:
    if not checkpoint_path:
        return
    try:
        tmp_path = checkpoint_path + ".tmp"
        os.makedirs(os.path.dirname(checkpoint_path), exist_ok=True)
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(payload or {}, f, ensure_ascii=False, indent=2)
        os.replace(tmp_path, checkpoint_path)
    except Exception as exc:
        logger.warning("Failed to save OCR checkpoint %s: %s", checkpoint_path, exc)


def _capture_run_format(paragraph):
    if not getattr(paragraph, "runs", None):
        return {}
    run = paragraph.runs[0]
    font = run.font
    size = None
    try:
        size = font.size.pt if font.size else None
    except Exception:
        size = None
    return {
        "bold": bool(run.bold) if run.bold is not None else None,
        "italic": bool(run.italic) if run.italic is not None else None,
        "underline": bool(run.underline) if run.underline is not None else None,
        "font_name": font.name,
        "font_size_pt": size,
        "style_name": getattr(getattr(run, "style", None), "name", None),
    }


def _build_docx_markdown_and_layout(docx_input):
    doc = Document(docx_input)
    blocks = []

    def _append_block(kind, block_id, text, **extra):
        clean = str(text or "").strip()
        if not clean:
            return
        blocks.append(
            {
                "block_id": block_id,
                "kind": kind,
                "text": clean,
                **extra,
            }
        )

    for para_idx, para in enumerate(doc.paragraphs):
        _append_block(
            "paragraph",
            f"paragraph:{para_idx}",
            para.text,
            paragraph_index=para_idx,
            style_name=getattr(getattr(para, "style", None), "name", None),
            run_format=_capture_run_format(para),
        )

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()
                first_para = cell.paragraphs[0] if cell.paragraphs else None
                _append_block(
                    "table_cell",
                    f"table:{table_idx}:{row_idx}:{cell_idx}",
                    cell_text,
                    table_index=table_idx,
                    row_index=row_idx,
                    cell_index=cell_idx,
                    style_name=getattr(getattr(first_para, "style", None), "name", None) if first_para else None,
                    run_format=_capture_run_format(first_para) if first_para else {},
                )

    for section_idx, section in enumerate(doc.sections):
        for para_idx, para in enumerate(section.header.paragraphs):
            _append_block(
                "header_paragraph",
                f"header:{section_idx}:{para_idx}",
                para.text,
                section_index=section_idx,
                paragraph_index=para_idx,
                style_name=getattr(getattr(para, "style", None), "name", None),
                run_format=_capture_run_format(para),
            )
        for para_idx, para in enumerate(section.footer.paragraphs):
            _append_block(
                "footer_paragraph",
                f"footer:{section_idx}:{para_idx}",
                para.text,
                section_index=section_idx,
                paragraph_index=para_idx,
                style_name=getattr(getattr(para, "style", None), "name", None),
                run_format=_capture_run_format(para),
            )

    markdown_parts = ["# Extracted Document", ""]
    for block in blocks:
        markdown_parts.append(f"<!--BLOCK:{block['block_id']}-->")
        markdown_parts.append(block["text"])
        markdown_parts.append("")

    return "\n".join(markdown_parts).strip(), {
        "version": 1,
        "source_format": "docx",
        "original_docx_path": os.path.abspath(docx_input),
        "blocks": blocks,
    }


def extract_text_from_docx(docx_input):
    markdown, _ = _build_docx_markdown_and_layout(docx_input)
    return markdown


def extract_text_from_doc(doc_input):
    result = subprocess.run(
        ["/usr/bin/textutil", "-convert", "txt", "-stdout", doc_input],
        check=True,
        capture_output=True,
        text=True,
    )
    return result.stdout


def extract_text_from_text_file(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _detect_extensionless_format(path: str) -> str:
    try:
        with open(path, "rb") as f:
            header = f.read(8)
    except Exception:
        return ""

    if header.startswith(b"%PDF-"):
        return "pdf"
    if header.startswith(b"PK"):
        try:
            with zipfile.ZipFile(path) as zf:
                names = set(zf.namelist())
            if "word/document.xml" in names or any(name.startswith("word/") for name in names):
                return "docx"
        except Exception:
            return ""
    return ""


def _guess_image_mime_type(path: str) -> str:
    ext = os.path.splitext(str(path or ""))[1].lower()
    if ext in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if ext == ".png":
        return "image/png"
    if ext == ".webp":
        return "image/webp"
    return "image/png"


def _default_vision_model(provider: str) -> str:
    provider_name = str(provider or "").strip().lower()
    if provider_name == "google_studio":
        return "gemma-4-31b-it"
    return "gemma4:e4b"


def _extract_image_with_google_model(path: str, vision_model: str, api_key: str | None = None) -> tuple[str, str]:
    with open(path, "rb") as f:
        image_bytes = f.read()
    key = str(api_key or "").strip() or os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY") or ""
    if not key:
        raise ValueError("Google API key is required for google_studio vision provider")
    client = genai.Client(api_key=key)
    prompt = (
        "你是法律文书图片OCR助手。请尽量忠实提取图片中可见的全部文字。"
        "只返回 JSON，格式为 {\"text\":\"完整文本\",\"summary\":\"一句话概述\"}。"
        "不要输出额外解释。"
    )
    response = client.models.generate_content(
        model=vision_model,
        contents=[
            prompt,
            types.Part.from_bytes(data=image_bytes, mime_type=_guess_image_mime_type(path)),
        ],
        config={"temperature": 0.0},
    )
    raw = str(getattr(response, "text", "") or "").strip()
    parsed = _extract_json_from_response(raw)
    return str((parsed or {}).get("text") or "").strip(), str((parsed or {}).get("summary") or "").strip()


def extract_text_from_image(
    path,
    vision_provider: str = "ollama",
    vision_model: str | None = None,
    vision_base_url: str = "http://127.0.0.1:11434",
) -> tuple[str, dict]:
    ocr_lines = _extract_image_lines_with_tesseract(path)
    tesseract_text = "\n".join(line["text"] for line in ocr_lines if str(line.get("text") or "").strip()).strip()

    provider = str(vision_provider or "ollama").strip().lower()
    resolved_vision_model = str(vision_model or _default_vision_model(provider)).strip()
    gemma_text = ""
    gemma_summary = ""
    try:
        if provider == "google_studio":
            gemma_text, gemma_summary = _extract_image_with_google_model(path, resolved_vision_model)
        else:
            with open(path, "rb") as f:
                image_b64 = base64.b64encode(f.read()).decode("utf-8")

            prompt = (
                "你是法律文书图片OCR助手。请尽量忠实提取图片中可见的全部文字。"
                "只返回 JSON，格式为 {\"text\":\"完整文本\",\"summary\":\"一句话概述\"}。"
                "不要输出额外解释。"
            )
            payload = {
                "model": resolved_vision_model,
                "prompt": prompt,
                "images": [image_b64],
                "stream": False,
            }

            with httpx.Client(timeout=120.0) as client:
                resp = client.post(f"{vision_base_url.rstrip('/')}/api/generate", json=payload)
                resp.raise_for_status()
                raw = str((resp.json() or {}).get("response") or "").strip()

            parsed = _extract_json_from_response(raw)
            gemma_text = str((parsed or {}).get("text") or "").strip()
            gemma_summary = str((parsed or {}).get("summary") or "").strip()
    except Exception as e:
        logger.warning("Gemma vision OCR fallback unavailable for image %s: %s", path, e)

    final_text = gemma_text or tesseract_text
    if not final_text:
        raise ValueError("No OCR text extracted from image")
    return final_text, {
        "vision_provider": vision_provider,
        "vision_model": resolved_vision_model,
        "vision_base_url": vision_base_url,
        "summary": gemma_summary,
        "ocr_engine": "tesseract+gemma" if str(vision_provider or "ollama").strip().lower() == "ollama" else "tesseract+google_studio",
        "ocr_lines": ocr_lines,
        "tesseract_text": tesseract_text,
        "vision_text": gemma_text,
    }


def extract_text(
    path,
    *,
    vision_provider: str = "ollama",
    vision_model: str | None = None,
    vision_base_url: str = "http://127.0.0.1:11434",
    checkpoint_path: str | None = None,
):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        text, pages_ocr_data = extract_text_from_pdf(
            path,
            vision_provider=vision_provider,
            vision_model=vision_model,
            vision_base_url=vision_base_url,
            checkpoint_path=checkpoint_path,
        )
        layout_map = build_pdf_layout_map(path, extracted_text=text, pages_ocr_data=pages_ocr_data)
        if checkpoint_path:
            layout_map["ocr_checkpoint_path"] = os.path.abspath(checkpoint_path)
        return text, "pdf", layout_map
    if ext in {".png", ".jpg", ".jpeg"}:
        text, meta = extract_text_from_image(
            path,
            vision_provider=vision_provider,
            vision_model=vision_model,
            vision_base_url=vision_base_url,
        )
        return text, "image", meta
    if ext == ".docx":
        markdown, layout_map = _build_docx_markdown_and_layout(path)
        return markdown, "docx", layout_map
    if ext == ".doc":
        return extract_text_from_doc(path), "doc", None
    if ext in {".md", ".markdown", ".txt"}:
        return extract_text_from_text_file(path), ext.lstrip("."), None
    if not ext:
        detected = _detect_extensionless_format(path)
        if detected == "pdf":
            text, pages_ocr_data = extract_text_from_pdf(
                path,
                vision_provider=vision_provider,
                vision_model=vision_model,
                vision_base_url=vision_base_url,
                checkpoint_path=checkpoint_path,
            )
            layout_map = build_pdf_layout_map(path, extracted_text=text, pages_ocr_data=pages_ocr_data)
            if checkpoint_path:
                layout_map["ocr_checkpoint_path"] = os.path.abspath(checkpoint_path)
            return text, "pdf", layout_map
        if detected == "docx":
            markdown, layout_map = _build_docx_markdown_and_layout(path)
            return markdown, "docx", layout_map
        return extract_text_from_text_file(path), "txt", None
    raise ValueError(f"Unsupported file type: {ext}")

def run_extractor(
    input_file,
    projects_root,
    optional_case_id=None,
    optional_mode="redact",
    *,
    vision_provider: str = "ollama",
    vision_model: str | None = None,
    vision_base_url: str = "http://127.0.0.1:11434",
    checkpoint_path: str | None = None,
):
    """
    Extracts text from a PDF file and formats it as Markdown.
    
    Args:
        input_file: Path to the input PDF file.
        projects_root: Root directory for projects.
        optional_case_id: Optional case identifier.
        optional_mode: Operation mode (redact / restore / log).
        
    Returns:
        Result dictionary with 'ok' status and 'markdown_content' or 'markdown_path'.
    """
    logger.info(f"Extracting text from {input_file}")
    
    if not os.path.exists(input_file):
        logger.error(f"PDF file not found: {input_file}")
        return {"ok": False, "error": f"PDF file not found: {input_file}"}
        
    try:
        if checkpoint_path is None and str(input_file or "").lower().endswith(".pdf"):
            checkpoint_path = _build_pdf_checkpoint_path(projects_root, input_file)
        text, source_format, layout_map = extract_text(
            input_file,
            vision_provider=vision_provider,
            vision_model=vision_model,
            vision_base_url=vision_base_url,
            checkpoint_path=checkpoint_path,
        )
        
        if not text.strip():
            logger.warning(f"No text extracted from {input_file}")
            return {"ok": False, "error": "No text extracted from input document. Ensure it is text-based."}
            
        # Format as Markdown
        if source_format == "docx" and str(text).lstrip().startswith("# Extracted Document"):
            markdown_content = text
        else:
            markdown_content = f"# Extracted Document\n\n{text}"
        
        return {
            "ok": True, 
            "markdown_content": markdown_content,
            "pdf_path": input_file,
            "source_format": source_format,
            "layout_map": layout_map,
        }
        
    except Exception as e:
        logger.exception(f"Error extracting text from {input_file}: {str(e)}")
        return {"ok": False, "error": f"Extraction error: {str(e)}"}

def _extract_json_from_response(raw: str) -> dict:
    text = str(raw or "").strip()
    match = re.search(r"(\{.*\})", text, re.DOTALL)
    candidate = match.group(1) if match else text
    try:
        parsed = json.loads(candidate)
        if isinstance(parsed, dict):
            return parsed
    except Exception:
        pass
    return {"text": text, "summary": ""}


def _extract_image_lines_with_tesseract(path: str, lang: str = "chi_sim+eng") -> list[dict]:
    try:
        result = subprocess.run(
            ["tesseract", path, "stdout", "-l", lang, "--psm", "6", "tsv"],
            check=True,
            capture_output=True,
            text=True,
        )
    except Exception as e:
        logger.warning("Tesseract OCR failed for %s: %s", path, e)
        return []

    rows = []
    lines: dict[tuple[int, int, int], dict] = {}
    raw_lines = result.stdout.splitlines()
    if not raw_lines:
        return rows
    header = raw_lines[0].split("\t")
    for raw in raw_lines[1:]:
        cols = raw.split("\t")
        if len(cols) != len(header):
            continue
        row = dict(zip(header, cols))
        text = str(row.get("text") or "").strip()
        if not text:
            continue
        try:
            level = int(row.get("level") or 0)
            left = int(float(row.get("left") or 0))
            top = int(float(row.get("top") or 0))
            width = int(float(row.get("width") or 0))
            height = int(float(row.get("height") or 0))
            block_num = int(row.get("block_num") or 0)
            par_num = int(row.get("par_num") or 0)
            line_num = int(row.get("line_num") or 0)
        except Exception:
            continue
        if level != 5:
            continue
        key = (block_num, par_num, line_num)
        entry = lines.setdefault(
            key,
            {
                "text_parts": [],
                "tokens": [],
                "left": left,
                "top": top,
                "right": left + width,
                "bottom": top + height,
            },
        )
        entry["text_parts"].append(text)
        entry["tokens"].append(
            {
                "text": text,
                "bbox": [left, top, left + width, top + height],
            }
        )
        entry["left"] = min(entry["left"], left)
        entry["top"] = min(entry["top"], top)
        entry["right"] = max(entry["right"], left + width)
        entry["bottom"] = max(entry["bottom"], top + height)

    for _, entry in sorted(lines.items(), key=lambda item: (item[0][0], item[0][1], item[0][2])):
        text = " ".join(entry["text_parts"]).strip()
        if not text:
            continue
        rows.append(
            {
                "text": text,
                "bbox": [entry["left"], entry["top"], entry["right"], entry["bottom"]],
                "tokens": entry.get("tokens", []),
            }
        )
    return rows


if __name__ == "__main__":
    # Simple CLI for testing
    import sys
    if len(sys.argv) > 1:
        res = run_extractor(sys.argv[1], ".")
        if res["ok"]:
            print(res["markdown_content"])
        else:
            print(f"Error: {res['error']}")
